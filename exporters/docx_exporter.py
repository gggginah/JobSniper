from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_COLOR_INDEX
from utils.file_io import read_file
from utils.markdown import clean_markdown_bold


def convert_markdown_to_docx(markdown_path, docx_path):
    document = Document()

    section = document.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)

    body_font = "Arial"
    body_size = Pt(9.3)

    normal_style = document.styles["Normal"]
    normal_style.font.name = body_font
    normal_style.font.size = body_size

    bullet_style = document.styles["List Bullet"]
    bullet_style.font.name = body_font
    bullet_style.font.size = body_size


    markdown = read_file(markdown_path)
    lines = markdown.splitlines()

    for line in lines:
        line = line.strip()

        if not line:
            continue

        if line.startswith("# "):

            name = line.replace("# ", "").strip()

            paragraph = document.add_paragraph()

            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            paragraph.paragraph_format.space_after = Pt(2)

            run = paragraph.add_run(name)

            run.bold = True
            run.font.name = "Arial"
            run.font.size = Pt(18)

        elif "|" in line and (
            "@" in line
            or "+86" in line
        ):

            paragraph = document.add_paragraph()

            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            paragraph.paragraph_format.space_after = Pt(4)

            run = paragraph.add_run(line)

            run.font.name = body_font
            run.font.size = Pt(8.5)
            
        elif line.startswith("## "):

            heading = line.replace("## ", "").strip().upper()

            paragraph = document.add_paragraph()

            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

            paragraph.paragraph_format.space_before = Pt(6)
            paragraph.paragraph_format.space_after = Pt(2)

            run = paragraph.add_run(heading)

            run.bold = True
            run.font.name = body_font
            run.font.size = Pt(11)

            add_bottom_border(paragraph)

        elif "|" in line:
    
            paragraph = document.add_paragraph()

            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

            run = paragraph.add_run(clean_markdown_bold(line))

            run.font.name = body_font
            run.font.size = body_size

        elif line.startswith("Project:"):

            paragraph = document.add_paragraph()

            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

            run = paragraph.add_run(
                clean_markdown_bold(line)
            )

            run.bold = True
            run.font.name = body_font
            run.font.size = body_size

        elif line.startswith("- "):
            bullet_text = line[2:].strip()
            paragraph = document.add_paragraph(style="List Bullet")
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0

            run = paragraph.add_run(clean_markdown_bold(bullet_text))
            run.font.name = body_font
            run.font.size = body_size


        else:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(2)

            if line.startswith("**") and "**" in line[2:]:
                clean_line = clean_markdown_bold(line)
                run = paragraph.add_run(clean_line)
                run.font.name = body_font
                run.font.size = body_size

                run.bold = True
            
            else:

                paragraph = document.add_paragraph()

                run = paragraph.add_run(
                    clean_markdown_bold(line)
                )

                run.font.name = body_font
                run.font.size = body_size


    document.save(docx_path)



def add_bottom_border(paragraph):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    p = paragraph._p
    pPr = p.get_or_add_pPr()

    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)

    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "0")
    bottom.set(qn("w:color"), "000000")

    pBdr.append(bottom)

