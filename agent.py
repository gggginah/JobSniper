from openai import OpenAI
from dotenv import load_dotenv
from datetime import datetime
import os
import time
import sys
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_COLOR_INDEX

load_dotenv()


client = OpenAI(
    api_key=os.getenv("DEEPSEEK_API_KEY"),
    base_url="https://api.deepseek.com"
)


JD_PATH = "JDs/jd.txt"
EXPERIENCE_BANK_PATH = "experience_bank.md"


RESUME_FILES = {
    "SE": "resumes/SupportEngineer.txt",
    "CSM": "resumes/CustomerSuccessManager.txt",
    "TAM": "resumes/TechnicalAccountManager.txt"
}


ROLE_TITLE_KEYWORDS = {
    "SE": [
        "solution engineer",
        "solutions engineer",
        "support engineer",
        "technical support engineer",
        "customer support engineer",
        "sales engineer",
        "pre-sales engineer",
        "presales engineer",
        "solution consultant",
        "solutions consultant"
    ],

    "PRESALES_SUPPORT": [
        "technical sales",
        "technical sales support",
        "technical sales support specialist",
        "sales support specialist",
        "technical sales specialist",
        "pre-sales",
        "presales",
        "solution support",
        "pre-sales support",
        "presales support",
        "application support specialist",
        "product support specialist",
        "technical specialist"
    ],

    "CSM": [
        "customer success manager",
        "customer success",
        "client success manager",
        "client success",
        "customer experience",
        "customer adoption",
        "account manager",
        "customer relationship manager"
    ],

    "TAM": [
        "technical account manager",
        "technical account",
        "tam",
        "implementation manager",
        "implementation consultant",
        "technical consultant",
        "customer engineer",
        "technical customer success manager"
    ]
}


ROUTE_TO_RESUME_PROFILE = {
    "SE": "SE",
    "PRESALES_SUPPORT": "SE",
    "CSM": "CSM",
    "TAM": "TAM"
}


FINAL_OUTPUT_MARKERS = [
    "===TAILORED_RESUME_EN_MD===",
    "===REVIEW_NOTES_BILINGUAL_MD==="
]


EVIDENCE_MARKERS = [
    "===SELECTED_EVIDENCE_MD==="
]


def read_file(path):
    with open(path, "r", encoding="utf-8") as file:
        return file.read()


def save_file(path, content):
    with open(path, "w", encoding="utf-8") as file:
        file.write(content)


def read_jd_file(path):
    content = read_file(path)
    lines = content.splitlines()

    if len(lines) == 0:
        raise ValueError("JD 文件是空的，请在 JDs/jd.txt 中填入岗位名称和 JD 内容。")

    job_title_line = lines[0].strip()
    job_title = job_title_line.replace("Job Title:", "").strip()
    jd_text = "\n".join(lines[1:]).strip()

    if not job_title:
        raise ValueError("没有读取到岗位名称，请把岗位名称放在 JDs/jd.txt 第一行。")

    if not jd_text:
        raise ValueError("没有读取到 JD 内容，请把完整 JD 放在岗位名称下面。")

    return job_title, jd_text


def load_all_resumes():
    resumes = {}

    for role, path in RESUME_FILES.items():
        resumes[role] = read_file(path)

    return resumes


def choose_resume_by_job_title(job_title):
    title = job_title.lower()
    scores = {}

    for role, keywords in ROLE_TITLE_KEYWORDS.items():
        score = 0

        for keyword in keywords:
            if keyword in title:
                score += 1

        scores[role] = score

    best_role = max(scores, key=scores.get)

    if scores[best_role] == 0:
        return "UNKNOWN", scores

    return best_role, scores


def get_stage2_folder_from_args():
    if "--stage2" in sys.argv:
        index = sys.argv.index("--stage2")

        if index + 1 >= len(sys.argv):
            raise ValueError("请在 --stage2 后面提供 output 文件夹路径。")

        return sys.argv[index + 1]

    return None


def sanitize_job_title(job_title):
    safe_job_title = (
        job_title
        .replace(" ", "_")
        .replace("/", "_")
        .replace("\\", "_")
        .replace(":", "")
        .replace("*", "")
        .replace("?", "")
        .replace('"', "")
        .replace("<", "")
        .replace(">", "")
        .replace("|", "")
    )

    return safe_job_title


def create_output_run_folder(job_title):
    timestamp = datetime.now().strftime("%Y-%m-%d_%H%M")
    safe_job_title = sanitize_job_title(job_title)

    folder_name = f"{timestamp}_{safe_job_title}"
    output_path = os.path.join("output", folder_name)

    os.makedirs(output_path, exist_ok=True)

    return output_path


def extract_markdown_section(full_text, start_marker, end_marker=None):
    if start_marker not in full_text:
        return ""

    start_index = full_text.find(start_marker) + len(start_marker)

    if end_marker and end_marker in full_text:
        end_index = full_text.find(end_marker)
        return full_text[start_index:end_index].strip()

    return full_text[start_index:].strip()


def ask_deepseek(prompt, required_markers=None, max_tokens=8000):
    max_retries = 3

    for attempt in range(max_retries):
        try:
            final_prompt = prompt

            if required_markers and attempt > 0:
                final_prompt = prompt + """

IMPORTANT:
Your previous response was incomplete or did not follow the required output format.
You must include all required exact markers.
Do not rename the markers.
Do not omit required sections.
"""

            response = client.chat.completions.create(
                model="deepseek-chat",
                messages=[
                    {"role": "user", "content": final_prompt}
                ],
                temperature=0,
                max_tokens=max_tokens,
                timeout=60
            )

            result = response.choices[0].message.content

            if required_markers:
                missing_markers = [
                    marker for marker in required_markers
                    if marker not in result
                ]

                if not missing_markers:
                    return result

                print(f"DeepSeek output format incomplete. Attempt {attempt + 1}/{max_retries}")
                print("Missing markers:", missing_markers)

                if attempt < max_retries - 1:
                    print("Retrying with stricter format instruction in 5 seconds...")
                    time.sleep(5)
                    continue

                raise ValueError(f"DeepSeek response missing required markers: {missing_markers}")

            return result

        except Exception as error:
            print(f"DeepSeek API call failed. Attempt {attempt + 1}/{max_retries}")
            print("Error:", error)

            if attempt < max_retries - 1:
                print("Retrying in 5 seconds...")
                time.sleep(5)
            else:
                raise


def build_fallback_resume_selection_prompt(job_title, jd, title_scores):
    prompt = f"""
You are a resume routing assistant.

The local title-based router could not confidently match this job title to a resume profile.

Your task is to analyze the full job description and select the most suitable resume profile.

Available resume profiles:

1. SE
Use this profile when the role is closest to:
- Support Engineer
- Technical Support Engineer
- Product Support
- Application Support
- Solution Engineer
- Pre-sales Engineer
- Technical troubleshooting
- Technical customer communication

2. PRESALES_SUPPORT
Use this profile when the role is closest to:
- Technical Sales
- Technical Sales Support
- Sales Support with technical responsibilities
- Solution support
- Pre-sales support
- Product consultation
- Customer-facing technical sales support

3. CSM
Use this profile when the role is closest to:
- Customer Success Manager
- Customer Success
- Customer Adoption
- Customer Relationship Management
- Account growth
- Renewal / retention support
- Customer onboarding from a relationship perspective

4. TAM
Use this profile when the role is closest to:
- Technical Account Manager
- Implementation Consultant
- Technical Consultant
- Customer Engineer
- Enterprise technical relationship management
- Post-sales technical ownership
- Implementation / integration support

5. UNKNOWN
Use UNKNOWN only when the JD is clearly outside these directions or there is not enough evidence to choose.

Important rules:
- Do not rely only on the job title.
- Use the full JD responsibilities and requirements.
- Prefer the profile that best matches the actual work described in the JD.
- If the role is hybrid, choose the closest primary profile.
- Return only one final selected profile.
- Do not generate a resume.
- Do not generate review notes.
- Keep the reason brief.

Output format:
You must use exactly this format:

SELECTED_PROFILE: SE / PRESALES_SUPPORT / CSM / TAM / UNKNOWN
REASON_EN: one short English explanation
REASON_ZH: 一句简短中文解释

====================
Job Title
====================
{job_title}

====================
Local Router Scores
====================
{title_scores}

====================
Job Description
====================
{jd}
"""
    return prompt


def parse_selected_resume_profile(selection_result):
    valid_profiles = ["SE", "PRESALES_SUPPORT", "CSM", "TAM", "UNKNOWN"]

    for line in selection_result.splitlines():
        line = line.strip()

        if line.startswith("SELECTED_PROFILE:"):
            profile = line.replace("SELECTED_PROFILE:", "").strip()

            if profile in valid_profiles:
                return profile

    return "UNKNOWN"


def build_evidence_selection_prompt(job_title, jd, experience_bank):
    prompt = f"""
You are an evidence selection assistant.

Your task is to read the job description and the candidate's Chinese Experience Bank,
then select only the most relevant evidence for this specific role.

The Experience Bank is written mainly in Chinese.
Chinese facts are the source of truth.

Important rules:
- Do not generate a resume.
- Do not rewrite the full Experience Bank.
- Do not translate the entire Experience Bank.
- Select only facts that are relevant to the JD.
- Keep the output concise and structured.
- Do not invent facts, metrics, tools, titles, customers, or responsibilities.
- Preserve project ownership strictly.
- Do not mix evidence between projects.
- If you are unsure which project a fact belongs to, put it under "Needs human review".
- If a fact is useful but risky or unclear, mark it as "Needs human review".

Project separation rules:
- Facts from Project 01 must stay under Project 01.
- Facts from Project 02 must stay under Project 02.
- JobSniper facts must stay under Personal Project.
- Do not move metrics from one project to another.
- Do not merge Project 01 and Project 02 into one generic experience section.
- Do not mix ECDS/NEBS evidence with the enterprise credit system reconstruction project.
- Do not mix discount business evidence with online acceptance business evidence.

Selection limits:
- Select no more than 8 facts per project.
- Select no more than 6 metrics in total.
- Select no more than 6 technical evidence items in total.
- Select no more than 6 support or user-facing evidence items in total.
- If JobSniper is not relevant to the JD, keep it brief or state "Not selected".

Output format:

===SELECTED_EVIDENCE_MD===

# Selected Evidence for This JD

## Job Target

- Job title:
- Main JD focus:
- Recommended resume emphasis:

---

## Project 01: 新一代企业授信用信系统重构 / 票据贴现业务

### Why relevant to this JD
-

### Selected facts
-

### Selected metrics
-

### Selected technical evidence
-

### Selected support / user-facing evidence
-

### Boundaries / do not overstate
-

### Needs human review
-

---

## Project 02: ECDS to NEBS / 线上承兑 / 线上汽车金融

### Why relevant to this JD
-

### Selected facts
-

### Selected metrics
-

### Selected technical evidence
-

### Selected support / user-facing evidence
-

### Boundaries / do not overstate
-

### Needs human review
-

---

## Personal Project: JobSniper AI Resume Agent

### Why relevant to this JD
-

### Selected facts
-

### Selected technical evidence
-

### Boundaries / do not overstate
-

### Needs human review
-

---

## Evidence Not Selected
-

====================
Job Title
====================
{job_title}

====================
Job Description
====================
{jd}

====================
Chinese Experience Bank
====================
{experience_bank}
"""
    return prompt


def build_single_resume_prompt(job_title, jd, selected_role, selected_resume, title_scores, selected_evidence):
    prompt = f"""
You are a senior recruiter, resume strategist, and ATS optimization assistant.

Your task is to generate a JD-tailored English resume draft and bilingual review notes.

You must use:
1. The selected resume as the base structure.
2. The job description as the target.
3. The selected evidence extracted from the candidate's Chinese Experience Bank.

Important language rules:
- The final resume must be written in English.
- The selected evidence may contain Chinese facts.
- Chinese facts are the source of truth.
- Translate only relevant Chinese facts into safe, natural, resume-ready English.
- Do not translate or use evidence that is irrelevant to this JD.
- Do not invent or embellish facts during translation.

Important resume generation rules:
- Use the selected resume as the base structure.
- Do not generate a resume completely from scratch.
- Preserve the original resume's project structure as much as possible.
- You may revise, reorder, replace, or strengthen bullet points only when supported by the selected resume or selected evidence.
- You may add new evidence from selected_evidence if it improves JD alignment.
- Do not add evidence that is not in the selected resume or selected evidence.
- Keep Project 01 and Project 02 separate.
- Do not merge bullets from different projects.
- Do not move metrics between projects.
- Do not move Project 01 metrics into Project 02.
- Do not move Project 02 metrics into Project 01.
- JobSniper evidence may only be used in a Personal Project section if relevant to the JD.
- If JobSniper is not relevant, do not force it into the resume.

Strict anti-hallucination rules:
- Do not invent facts, metrics, tools, titles, certifications, customers, platforms, or responsibilities.
- Do not claim formal CSAT, SLA, Contact Rate, revenue ownership, or official Product Manager ownership unless clearly supported.
- Do not upgrade internal users into external customers unless the evidence clearly supports it.
- Do not describe branch client managers as external customers.
- Do not claim production deployment or commercial users for JobSniper.
- If a selected evidence item is risky, unclear, or marked as "Needs human review", do not put it in the final resume. Mention it in review_notes_bilingual.md instead.
- The review_notes_bilingual.md section must be bilingual. If you provide an English point, also provide a Chinese explanation.

Scoring rules:
Use this fixed 100-point rubric:
- Must-have Requirements Match: 35
- Transferable Skills Match: 20
- Responsibility Alignment: 20
- Nice-to-have Coverage: 10
- ATS Readability: 15

Application decision rules:
After scoring, provide:
- Score Diagnosis
- Apply Priority: High / Medium / Low
- Stretch Fit Potential: High / Medium / Low
- Resume Fixability: High / Medium / Low
- Recommended Action: Apply / Apply with light tailoring / Apply with heavy tailoring / Do not prioritize

Output rules:
- You must output exactly two sections.
- Use the exact markers below.
- Do not rename the markers.
- Do not omit either section.
- Do not include extra text before the first marker.

Required output format:

===TAILORED_RESUME_EN_MD===

Write the full tailored English resume here.
The resume should be clean, ATS-friendly, and ready to copy into a Markdown file.
Do not include Chinese in this section.
Do not include review notes in this section.
Do not include risk comments in this section.

Markdown formatting requirements for ===TAILORED_RESUME_EN_MD===:

The final resume must use clean resume-style Markdown that can be converted into DOCX.

Required structure:
- Use "# Candidate Name" for the candidate name.
- Put the contact information on the line immediately after the candidate name.
- Use "##" for major resume sections, such as:
  ## PROFESSIONAL SUMMARY
  ## WORK EXPERIENCE
  ## PERSONAL PROJECT
  ## EDUCATION
  ## TECHNICAL SKILLS
- Use "**bold**" for company names, project names, school names, and personal project names.
- Use "- " for all bullet points.
- Use " | " to separate date ranges, locations, and other same-line metadata.
- For education labels such as 211 or Double 1st-Class, write them as "tag:211" and "tag:Double 1st-Class".
- Do not use tables.
- Do not use HTML.
- Do not include horizontal lines in Markdown.
- Do not include Markdown code blocks.

Example format:

# Xiaofei Han

+86 xxx | email@example.com | Beijing | GitHub URL

## PROFESSIONAL SUMMARY

One concise paragraph.

## WORK EXPERIENCE

**China CITIC Bank Co., Ltd** | Sep 2020 - Jan 2026
Technical Support & Integration Engineer | Beijing

**Project: Project Name**
- Bullet point
- Bullet point

## PERSONAL PROJECT

**JobSniper — LLM-powered Job Application Workflow Assistant** | Jun 2026 - Present
- Bullet point
- Bullet point

## EDUCATION

**Beijing University of Technology** | tag:211 | tag:Double 1st-Class | Aug 2016 - Jun 2020
Software Engineering Bachelor Full-time | Beijing

## TECHNICAL SKILLS

Programming & AI: Java, Python, LLM API Integration, Prompt Engineering, LLM Workflow Design, Git/GitHub
Integration & Troubleshooting: RESTful APIs, JSON, SQL, MySQL, Linux, Shell scripting, Log Analysis, Postman
Languages: English, Chinese


===REVIEW_NOTES_BILINGUAL_MD===

Write bilingual review notes here.

Strict bilingual rules:
- This section must include both English and Chinese.
- Every major heading must be bilingual, using this format:
  ## English Heading / 中文标题
- For each important point, write the English explanation first, then the Chinese explanation immediately below it.
- Do not write this section in English only.
- Do not write this section in Chinese only.
- Keep the final resume section English-only, but this review notes section must be bilingual.

Must include these bilingual sections:
1. Match Score and Rubric Breakdown / 匹配度评分与评分拆解
2. Score Diagnosis / 分数诊断
3. Apply Priority / 投递优先级
4. Stretch Fit Potential / Stretch 岗位潜力
5. Resume Fixability / 简历可修复度
6. Recommended Action / 建议行动
7. JD Keywords Used / 使用到的 JD 关键词
8. Evidence Added from Experience Bank / 从 Experience Bank 添加的证据
9. Evidence Not Used and Why / 未使用的证据及原因
10. Risky / Needs Human Review / 有风险或需要人工确认的内容
11. Project Separation Check / 项目归属检查
12. Potential Interview Questions and Suggested Answers / 潜在面试问题与建议回答

====================
Job Title
====================
{job_title}

====================
Selected Role Route
====================
{selected_role}

====================
Local Router Scores
====================
{title_scores}

====================
Job Description
====================
{jd}

====================
Selected Resume Base
====================
{selected_resume}

====================
Selected Evidence from Experience Bank
====================
{selected_evidence}
"""
    return prompt


def save_selected_evidence(selected_evidence, output_folder):
    save_file(
        os.path.join(output_folder, "selected_evidence.md"),
        selected_evidence
    )

    print(f"selected_evidence.md saved to: {output_folder}")


def save_final_outputs(result, output_folder):
    tailored_resume_en = extract_markdown_section(
        result,
        "===TAILORED_RESUME_EN_MD===",
        "===REVIEW_NOTES_BILINGUAL_MD==="
    )

    review_notes_bilingual = extract_markdown_section(
        result,
        "===REVIEW_NOTES_BILINGUAL_MD==="
    )

    tailored_resume_path = os.path.join(output_folder, "tailored_resume_en.md")
    review_notes_path = os.path.join(output_folder, "review_notes_bilingual.md")
    docx_path = os.path.join(output_folder, "tailored_resume.docx")

    if tailored_resume_en:
        save_file(tailored_resume_path, tailored_resume_en)
        print(f"tailored_resume_en.md saved to: {output_folder}")

        convert_markdown_to_docx(
            markdown_path=tailored_resume_path,
            docx_path=docx_path
        )
        print(f"tailored_resume.docx saved to: {output_folder}")

    else:
        print("Warning: tailored_resume_en.md was not generated. Resume marker not found.")

    if review_notes_bilingual:
        save_file(review_notes_path, review_notes_bilingual)
        print(f"review_notes_bilingual.md saved to: {output_folder}")
    else:
        print("Warning: review_notes_bilingual.md was not generated. Review notes marker not found.")

    print(f"Final files saved to: {output_folder}")



def select_resume(job_title, jd):
    print("开始进行本地简历路由...")

    best_role, title_scores = choose_resume_by_job_title(job_title)

    print("本地路由分数:", title_scores)

    resumes = load_all_resumes()

    if best_role == "UNKNOWN":
        print("岗位名称未匹配到明确简历方向，启动 fallback：让 DeepSeek 根据完整 JD 选择简历方向。")

        selection_prompt = build_fallback_resume_selection_prompt(
            job_title=job_title,
            jd=jd,
            title_scores=title_scores
        )

        selection_result = ask_deepseek(selection_prompt)

        print("\n====================")
        print("Fallback Resume Selection Result")
        print("====================\n")
        print(selection_result)

        selected_route = parse_selected_resume_profile(selection_result)

        if selected_route == "UNKNOWN":
            raise ValueError("Fallback could not select a suitable resume profile. Please review the JD manually.")

        selected_resume_profile = ROUTE_TO_RESUME_PROFILE[selected_route]
        selected_resume = resumes[selected_resume_profile]

        print(f"Fallback 选择岗位方向：{selected_route}")
        print(f"实际使用简历：{selected_resume_profile}")

        return selected_route, selected_resume, title_scores

    selected_resume_profile = ROUTE_TO_RESUME_PROFILE.get(best_role, best_role)
    selected_resume = resumes[selected_resume_profile]

    print(f"本地路由成功：岗位方向为 {best_role}，实际使用 {selected_resume_profile} 简历。")

    return best_role, selected_resume, title_scores


def run_stage1_evidence_selection(job_title, jd, output_folder):
    print("开始 Stage 1：根据 JD 从 Experience Bank 选择相关 evidence...")

    if not os.path.exists(EXPERIENCE_BANK_PATH):
        raise FileNotFoundError("没有找到 experience_bank.md，请先在项目根目录创建该文件。")

    experience_bank = read_file(EXPERIENCE_BANK_PATH)

    evidence_prompt = build_evidence_selection_prompt(
        job_title=job_title,
        jd=jd,
        experience_bank=experience_bank
    )

    selected_evidence_result = ask_deepseek(
        evidence_prompt,
        required_markers=EVIDENCE_MARKERS,
        max_tokens=5000
    )

    selected_evidence = extract_markdown_section(
        selected_evidence_result,
        "===SELECTED_EVIDENCE_MD==="
    )

    if not selected_evidence:
        raise ValueError("Stage 1 failed: selected evidence was empty.")

    save_selected_evidence(selected_evidence, output_folder)

    return selected_evidence


def run_stage2_resume_generation(job_title, jd, selected_role, selected_resume, title_scores, selected_evidence, output_folder):
    print("开始 Stage 2：基于 selected resume + selected evidence 生成最终简历...")

    prompt = build_single_resume_prompt(
        job_title=job_title,
        jd=jd,
        selected_role=selected_role,
        selected_resume=selected_resume,
        title_scores=title_scores,
        selected_evidence=selected_evidence
    )

    result = ask_deepseek(
        prompt,
        required_markers=FINAL_OUTPUT_MARKERS,
        max_tokens=9000
    )

    print("\n====================")
    print("DeepSeek Final Result")
    print("====================\n")
    print(result)

    save_final_outputs(result, output_folder)


def add_bottom_border(paragraph):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    p = paragraph._p
    pPr = p.get_or_add_pPr()

    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)

    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "0")
    bottom.set(qn("w:color"), "000000")

    pBdr.append(bottom)
    
def convert_markdown_to_docx(markdown_path, docx_path):
    document = Document()

    section = document.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)

    body_font = "Arial"
    body_size = Pt(9.3)

    normal_style = document.styles["Normal"]
    normal_style.font.name = body_font
    normal_style.font.size = body_size

    bullet_style = document.styles["List Bullet"]
    bullet_style.font.name = body_font
    bullet_style.font.size = body_size


    markdown = read_file(markdown_path)
    lines = markdown.splitlines()

    for line in lines:
        line = line.strip()

        if not line:
            continue

        if line.startswith("# "):

            name = line.replace("# ", "").strip()

            paragraph = document.add_paragraph()

            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            paragraph.paragraph_format.space_after = Pt(2)

            run = paragraph.add_run(name)

            run.bold = True
            run.font.name = "Arial"
            run.font.size = Pt(18)

        elif "|" in line and (
            "@" in line
            or "+86" in line
        ):

            paragraph = document.add_paragraph()

            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            paragraph.paragraph_format.space_after = Pt(4)

            run = paragraph.add_run(line)

            run.font.name = body_font
            run.font.size = Pt(8.5)
            
        elif line.startswith("## "):

            heading = line.replace("## ", "").strip().upper()

            paragraph = document.add_paragraph()

            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

            paragraph.paragraph_format.space_before = Pt(6)
            paragraph.paragraph_format.space_after = Pt(2)

            run = paragraph.add_run(heading)

            run.bold = True
            run.font.name = body_font
            run.font.size = Pt(11)

            add_bottom_border(paragraph)

        elif "|" in line:
    
            paragraph = document.add_paragraph()

            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

            run = paragraph.add_run(clean_markdown_bold(line))

            run.font.name = body_font
            run.font.size = body_size

        elif line.startswith("Project:"):

            paragraph = document.add_paragraph()

            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

            run = paragraph.add_run(
                clean_markdown_bold(line)
            )

            run.bold = True
            run.font.name = body_font
            run.font.size = body_size

        elif line.startswith("- "):
            bullet_text = line[2:].strip()
            paragraph = document.add_paragraph(style="List Bullet")
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0

            run = paragraph.add_run(clean_markdown_bold(bullet_text))
            run.font.name = body_font
            run.font.size = body_size


        else:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(2)

            if line.startswith("**") and "**" in line[2:]:
                clean_line = clean_markdown_bold(line)
                run = paragraph.add_run(clean_line)
                run.font.name = body_font
                run.font.size = body_size

                run.bold = True
            
            else:

                paragraph = document.add_paragraph()

                run = paragraph.add_run(
                    clean_markdown_bold(line)
                )

                run.font.name = body_font
                run.font.size = body_size


    document.save(docx_path)


def clean_markdown_bold(text):

    replacements = {
        "tag:211": "211 University",
        "tag:Double 1st-Class": "Double First-Class"
    }

    for old, new in replacements.items():
        text = text.replace(old, new)

    return text.replace("**", "")


def get_export_docx_folder_from_args():
    if "--export-docx" in sys.argv:
        index = sys.argv.index("--export-docx")

        if index + 1 >= len(sys.argv):
            raise ValueError("请在 --export-docx 后面提供 output 文件夹路径。")

        return sys.argv[index + 1]

    return None


def main():

    export_docx_folder = get_export_docx_folder_from_args()

    if export_docx_folder:

        markdown_path = os.path.join(export_docx_folder, "tailored_resume_en.md")

        docx_path = os.path.join(export_docx_folder, "tailored_resume.docx")

        if not os.path.exists(markdown_path):

            raise FileNotFoundError(f"没有找到 tailored_resume_en.md: {markdown_path}")

        convert_markdown_to_docx(

            markdown_path=markdown_path,

            docx_path=docx_path

        )

        print(f"tailored_resume.docx exported to: {export_docx_folder}")

        return

    job_title, jd = read_jd_file(JD_PATH)

    print("读取到岗位名称:", job_title)

    stage2_folder = get_stage2_folder_from_args()

    selected_role, selected_resume, title_scores = select_resume(job_title, jd)

    if stage2_folder:
        print("\n====================")
        print("Stage 2 only mode")
        print("====================")
        print(f"读取人工编辑后的 selected_evidence.md: {stage2_folder}")

        output_folder = stage2_folder
        selected_evidence_path = os.path.join(output_folder, "selected_evidence.md")

        if not os.path.exists(selected_evidence_path):
            raise FileNotFoundError(f"没有找到 selected_evidence.md: {selected_evidence_path}")

        selected_evidence = read_file(selected_evidence_path)

        if not selected_evidence.strip():
            raise ValueError("selected_evidence.md 是空的，请先补充内容。")

    else:
        print("\n====================")
        print("Full mode")
        print("====================")
        print("将先生成 selected_evidence.md，再生成最终简历。")

        output_folder = create_output_run_folder(job_title)

        selected_evidence = run_stage1_evidence_selection(
            job_title=job_title,
            jd=jd,
            output_folder=output_folder
        )

    run_stage2_resume_generation(
        job_title=job_title,
        jd=jd,
        selected_role=selected_role,
        selected_resume=selected_resume,
        title_scores=title_scores,
        selected_evidence=selected_evidence,
        output_folder=output_folder
    )


if __name__ == "__main__":
    main()